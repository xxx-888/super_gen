"""
Document Parser - 文档解析服务

支持从上传的文件中提取纯文本内容，用于剧本导入。
支持格式：.txt / .docx / .pdf
"""
import io
from typing import Tuple


def parse_document(file_bytes: bytes, filename: str) -> Tuple[str, str]:
    """从上传的文件中提取文本。

    Args:
        file_bytes: 文件二进制内容
        filename: 原始文件名（用于判断格式）

    Returns:
        (title, content) — title 是去掉扩展名的文件名，content 是提取的纯文本

    Raises:
        ValueError: 不支持的文件格式
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    # 标题：去掉扩展名
    title = filename.rsplit(".", 1)[0] if "." in filename else filename

    if ext == "txt":
        content = _parse_txt(file_bytes)
    elif ext == "docx":
        content = _parse_docx(file_bytes)
    elif ext == "pdf":
        content = _parse_pdf(file_bytes)
    else:
        raise ValueError(f"不支持的文件格式：.{ext}，请上传 .txt / .docx / .pdf 文件")

    # 清理：去掉首尾空白，多个连续空行压缩为单个
    lines = [line.rstrip() for line in content.splitlines()]
    cleaned: list[str] = []
    blank_streak = 0
    for line in lines:
        if line.strip() == "":
            blank_streak += 1
            if blank_streak <= 1:
                cleaned.append("")
        else:
            blank_streak = 0
            cleaned.append(line)
    content = "\n".join(cleaned).strip()

    return title, content


def _parse_txt(file_bytes: bytes) -> str:
    """纯文本文件：尝试 UTF-8，回退 GBK（兼容中文 Windows 文件）。"""
    for encoding in ("utf-8", "gbk", "gb2312", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return file_bytes.decode("utf-8", errors="replace")


def _parse_docx(file_bytes: bytes) -> str:
    """Word .docx 文档：用 python-docx 提取段落文本。"""
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                paragraphs.append("\t".join(row_texts))
    return "\n".join(paragraphs)


def _parse_pdf(file_bytes: bytes) -> str:
    """PDF 文档：用 pypdf 逐页提取文本。"""
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text.strip())
    return "\n\n".join(pages)
